# -*- coding: utf-8 -*-
"""One-off: generate mentor briefing Word doc from structured content."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parents[1] / "project-docs" / "현자-미팅-서비스-설명.docx"


def set_korean_font(run, name="맑은 고딕", size=11):
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_para(doc, text, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_korean_font(run, size=size)
    run.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    set_korean_font(run, size=16 if level == 1 else 14 if level == 2 else 12)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return h


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                set_korean_font(r, size=10)
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            table.rows[ri + 1].cells[ci].text = cell
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    set_korean_font(r, size=10)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_korean_font(run, size=11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_korean_font(run, size=11)


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("인디야(REDPEN)\n현자 미팅용 서비스 설명")
    tr.bold = True
    set_korean_font(tr, size=18)

    meta = [
        "작성일: 2026-06-01",
        "작성자: 루 (비개발자, 1인 제작)",
        "배포 URL: https://indiya-momo.github.io/REDPEN/",
        "버전: v0.1.0 전후 · 비공개/오픈 베타 준비 단계",
    ]
    for m in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(m)
        set_korean_font(run, size=10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    add_heading(doc, "1. 서비스가 뭔지 (한 줄)", 1)
    quote = doc.add_paragraph()
    qr = quote.add_run(
        "AI가 글을 고쳐 주는 앱이 아니라, 인디자인 등에서 만든 조판 PDF에서 "
        "맞춤법·표기 일관성을 규칙으로 자동 탐지하고, PDF 위에 위치를 표시하는 브라우저 로컬 도구입니다."
    )
    set_korean_font(qr)
    qr.italic = True

    add_para(doc, "포지셔닝: 맞춤법 앱 대체 X → 「PDF 교열·교정 전처리 스캐너」", bold=True)
    add_table(
        doc,
        ["원칙", "내용"],
        [
            ["자동 수정", "❌"],
            ["자동 탐지 + 하이라이트", "⭕"],
            ["원고·PDF 서버 업로드", "❌ (검수 본문은 브라우저 안에서만)"],
            ["AI 문장 추천·문맥 교정", "❌ (MVP 비목표)"],
        ],
    )
    add_para(doc, "브랜드: 대외 이름 「인디야: 교정냥 모모의 여행」 · 저장소/배포명 REDPEN")

    add_heading(doc, "2. 누가 쓰는지", 1)
    add_table(
        doc,
        ["대상", "Pain Point"],
        [
            ["출판사 편집자·교정·교열자", "같은 규칙·같은 오탈자를 PDF에서 반복 검수하는 피로"],
            ["인디자인 텍스트 PDF 제작자", "조판 PDF 단계에서 표기 통일·맞춤법 후보를 빠르게 훑고 싶음"],
            ["번역·기획 편집", "원고 유출 우려 → 서버에 PDF를 올리기 싫음"],
        ],
    )
    add_para(doc, "1차 타깃: PC · Chrome/Edge · 텍스트 PDF (스캔 PDF 불가)")
    add_para(doc, "모바일: 대문·모모의 방 등 미리보기 수준 — 본격 검수는 PC 권장")
    add_para(doc, "제작: 비개발자 1인 + AI(Cursor) 보조 개발 · 규칙 데이터는 Google 시트에서 관리 후 앱에 동기화")

    add_heading(doc, "3. 현재 되는 기능 (2026-06 기준)", 1)

    add_heading(doc, "사용자 흐름", 2)
    add_numbered(
        doc,
        [
            "대문(웰컴) → 검수 시작",
            "PDF 열기 (파일 선택 / 드래그 / 「PDF 열기」로 재연결)",
            "맞춤법 확인 또는 일관성 확인 탭",
            "적용할 규칙 ON/OFF",
            "검수 실행 → 좌측 결과 목록 + 우측 PDF 하이라이트",
            "항목 클릭 → 해당 페이지·위치로 이동",
            "피드백 보내기 (Google Form 또는 클립보드)",
        ],
    )

    add_heading(doc, "맞춤법·주의", 2)
    add_bullets(
        doc,
        [
            "내장 맞춤법 규칙 (시트 spelling_rules → JSON 동기화)",
            "주의(편집자 검토) 규칙 — find/replace·팁·켜기/끄기",
            "본보조(보조용언) 규칙 — 별도 시트, 할당량·검사에서 제외 등 정책 적용",
            "조판 PDF 줄바꿈·공백을 고려한 유연 매칭 (find → regex)",
            "규칙별 결과 묶음 · 진행률 표시 · 검사 중 버튼 비활성",
        ],
    )

    add_heading(doc, "일관성", 2)
    add_bullets(
        doc,
        [
            "사용자 일관성 규칙 등록 (표기 A/B 통일 등)",
            "전역 제외 구문",
            "복합·띄어쓰기 관련 패턴 일부",
        ],
    )

    add_heading(doc, "PDF·안정성", 2)
    add_bullets(
        doc,
        [
            "PDF.js로 텍스트 추출 · TextLayer 기반 하이라이트",
            "한글 cMap · 재저장 PDF 안내 (차단 대신 advisory)",
            "50MB 초과 → 검수 실행 불가 + PDF 나누기 외부 링크 (iLovePDF)",
            "스캔 PDF → 명확 거절 메시지",
            "IndexedDB 세션 복원 시도 (브라우저·권한에 따라 실패 가능)",
            "인쇄면 표시 설정 (스프레드 등)",
        ],
    )

    add_heading(doc, "UI·부가", 2)
    add_bullets(
        doc,
        [
            "대문 PC / 모바일 분리",
            "모모의 방 (비밀 공간, 방명록 등)",
            "버전 뱃지 (SemVer + 빌드 ID)",
            "PostHog 익명 통계 (opt-out 가능) · 피드백 모달",
        ],
    )

    add_heading(doc, "개발·품질 (운영자 관점)", 2)
    add_bullets(
        doc,
        [
            "테스트 108개 · husky pre-commit/pre-push",
            "GitHub Pages 자동 배포 (main push → docs/)",
            "규칙 JSON 스키마 검증 스크립트",
        ],
    )

    add_heading(doc, "아직 UI에 없거나 「준비 중」인 것 (중요)", 2)
    add_table(
        doc,
        ["항목", "실제"],
        [
            ["「기준 저장 준비 중」 문구", "규칙 세트 이름·여러 세트·저장 버튼 UI 미노출 (로직은 있음)"],
            ["규칙 세트 CRUD", "코드·localStorage는 있으나 검수 화면에서 숨김"],
            ["목차 검사", "UI·카피 플레이스홀더 수준 (본격 검사 미완)"],
            ["슬롯·1000개 한도", "코드에 남아 있으나 베타 UI에서는 슬롯 개념 비노출"],
            ["로그인·협업·클라우드 원고", "없음"],
        ],
    )

    add_heading(doc, "4. 앞으로 넣고 싶은 기능 (우선순위 감각)", 1)

    add_heading(doc, "베타 직전·직후 (현실적)", 2)
    add_bullets(
        doc,
        [
            "비공개/오픈 베타 공식화 — 태그 v0.1.0-beta.1, 테스터 가이드 공유",
            "배포 URL smoke 고정 · Known Issues 정리",
            "PostHog·피드백 Form 운영 검증",
            "50MB·300p 한도 — 베타 로그 보고 조정 여부 결정",
            "PDF 나누기 링크 등 안내 문구 정리 (외부 도구 연계 유지)",
        ],
    )

    add_heading(doc, "v1.1 이후 (기능·성능)", 2)
    add_bullets(
        doc,
        [
            "대용량·다규칙 — Web Worker, 청크 검사 고도화",
            "규칙 세트 UI 복구 (이름·복제·저장·피드백 UX)",
            "목차 검사 실제 구현 (장 제목 ↔ 본문 쪽 제목)",
            "MAX_RULES 한도 완화 또는 UI 정리",
            "CSS·번들 정리 (pdf.js lazy 등)",
            "앱 내 PDF 분할 — 난이도 높음, 당분간 외부 링크로 대체",
        ],
    )

    add_heading(doc, "하지 않을 것 (합의)", 2)
    add_bullets(
        doc,
        [
            "AI 자동 교정·문장 추천",
            "PDF 자동 수정·재생성·저장",
            "HWP · OCR · 스캔 PDF",
            "서버에 원고·검사 결과 저장",
            "「앞 100페이지만 검사」(사용자 혼란으로 채택 안 함)",
        ],
    )

    add_heading(doc, "5. 개발자(현자)에게 꼭 물어볼 점", 1)

    sections = [
        (
            "베타 go / no-go",
            [
                "지금 상태로 비공개 베타에 나가도 되는 최소 조건은 무엇인가? (smoke 항목 5개 정도로 줄여 달라)",
                "freeze 기간에 절대 손대지 말아야 할 영역 (ruleEngine, session, PDF gate 등) 확인.",
                "hotfix만 허용할 때, 어떤 변경은 hotfix가 아니라 다음 마일스톤인가?",
            ],
        ),
        (
            "제품·문서",
            [
                "PRD/와이어프레임 없이 키운 프로젝트인데, 지금 시점에 고정해야 할 문서는 무엇인가?",
                "테스터에게 「하는 일 / 하지 않는 일」 외에 반드시 읽게 할 문단 3줄은?",
                "50MB·300p 운영 한도 — 베타에서 유지 vs 완화, 데이터 없을 때 기본값 조언.",
            ],
        ),
        (
            "기술·구조 (1인 + AI 개발)",
            [
                "PC / 모바일 코드 분리 + git worktree 구조 — 장기적으로 괜찮은지, merge 전략.",
                "규칙 autosave는 되는데 UI는 「준비 중」 — 베타에서 오해를 막는 카피 한 줄 추천.",
                "IndexedDB 세션 복원 실패 — 베타 Known Issue로 두기 vs 반드시 고치기.",
                "PostHog 6종 이벤트 — 베타에 필요한 최소인지, 과한지.",
            ],
        ),
        (
            "운영·법·신뢰",
            [
                "피드백 Form에 원고 PDF 첨부 금지 — 안내 문구·NDA 관점에서 충분한지.",
                "통계 opt-out 문구 — 베타에서 법적·신뢰 측면 보완할 것.",
                "초대형 베타 몇 명·어떤 시나리오로 시작하는 게 좋은지 (편집자 2~3명 등).",
            ],
        ),
        (
            "로드맵 판단",
            [
                "목차 검사·규칙 세트 UI·Worker 중 베타 후 첫 1개로 뭘 추천하는지.",
                "인앱 PDF 나누기 vs 외부 링크 — 1인 프로젝트에서 언제 in-app을 검토할지.",
            ],
        ),
    ]
    qnum = 1
    for sub, qs in sections:
        add_heading(doc, sub, 2)
        for q in qs:
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(f"{qnum}. {q}")
            set_korean_font(run)
            qnum += 1

    add_heading(doc, "6. 참고 자료 (리포지토리 내)", 1)
    add_table(
        doc,
        ["문서", "용도"],
        [
            ["project-docs/product-spine.md", "제품 뼈대·로드맵·비목표"],
            ["project-docs/private-beta-guide.md", "테스터용 (하는 일/않는 일·FAQ)"],
            ["project-docs/pre-beta-3day-plan.md", "베타 전 체크리스트"],
            ["project-docs/analytics-beta.md", "PostHog"],
            ["project-docs/feedback-form.md", "피드백"],
        ],
    )

    add_heading(doc, "7. 미팅 시 한 장 요약 (구두용)", 1)
    add_bullets(
        doc,
        [
            "뭐: 조판 PDF 탐지 + 위치 표시, 로컬, AI 교정 아님",
            "누구: 편집·교정 현장, PDF 유출 싫은 사람",
            "지금: MVP 동작, 베타 준비, UI 일부 미완(규칙 세트·목차)",
            "다음: 베타 freeze → 피드백 → 성능/목차/규칙 UI",
            "물어볼 것: go/no-go, freeze 범위, 문서 최소선, 한도·세션·오해 문구",
        ],
    )

    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run(
        "이 문서는 현자 미팅용 초안입니다. 미팅 후 답변을 §5 아래에 「현자 답변」으로 붙이면 다음 액션 정리에 씁니다."
    )
    set_korean_font(fr, size=9)
    fr.italic = True
    fr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    main()
